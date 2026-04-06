#!/usr/bin/env python3
"""
Generate a Word (.docx) document for a PhD dissertation literature review chapter.

Usage:
    python generate_docx.py --input review.md --output literature_review.docx
    python generate_docx.py --input review.md --output literature_review.docx --title "选题名称" --chapter-num 2

Input: A markdown file with the literature review content. Expected structure:
    - Lines starting with ## are section headings (e.g., ## 2.1 核心概念界定)
    - Lines starting with ### are subsection headings
    - Lines starting with #### are sub-subsection headings
    - Other lines are body paragraphs
    - A blank line separates paragraphs
    - A section titled "参考文献" or "References" triggers reference formatting

Output: A formatted .docx file following Chinese PhD dissertation standards:
    - Body: SimSun (宋体) 12pt, 1.5x line spacing, first-line indent 2 chars
    - Chapter title: SimHei (黑体) 16pt, centered
    - Section titles: SimHei (黑体) 14pt, left-aligned
    - Subsection titles: SimHei (黑体) 12pt, left-aligned
    - References: hanging indent 1.27cm, SimSun 10.5pt
    - Margins: top/bottom 2.54cm, left 3.17cm, right 3.17cm
"""

import argparse
import re
import sys

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def set_normal_style(doc):
    """Configure the Normal style for body text (PhD dissertation standard)."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)  # 小四号
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.first_line_indent = Cm(0.74)  # 约2字符缩进
    # Set East Asian font to SimSun (宋体)
    rpr = style.element.find(qn("w:rPr"))
    if rpr is None:
        rpr = style.element.makeelement(qn("w:rPr"), {})
        style.element.append(rpr)
    ea = rpr.find(qn("w:rFonts"))
    if ea is None:
        ea = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(ea)
    ea.set(qn("w:eastAsia"), "宋体")


def set_east_asian_font(run, font_name):
    """Set East Asian font for a run element."""
    rpr = run._element.find(qn("w:rPr"))
    if rpr is None:
        rpr = run._element.makeelement(qn("w:rPr"), {})
        run._element.append(rpr)
    ea = rpr.find(qn("w:rFonts"))
    if ea is None:
        ea = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(ea)
    ea.set(qn("w:eastAsia"), font_name)


def add_chapter_title(doc, text):
    """Add chapter title: SimHei 16pt (三号), centered."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(24)
    para.paragraph_format.space_after = Pt(18)
    para.paragraph_format.first_line_indent = Cm(0)
    para.paragraph_format.line_spacing = 1.5
    run = para.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)  # 三号
    run.font.color.rgb = RGBColor(0, 0, 0)
    set_east_asian_font(run, "黑体")
    return para


def add_section_heading(doc, text, level):
    """Add section/subsection heading with proper formatting."""
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Cm(0)
    para.paragraph_format.line_spacing = 1.5

    if level == 2:
        # Section heading: SimHei 14pt (四号), left-aligned
        para.paragraph_format.space_before = Pt(18)
        para.paragraph_format.space_after = Pt(12)
        run = para.add_run(text)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)  # 四号
        run.font.color.rgb = RGBColor(0, 0, 0)
        set_east_asian_font(run, "黑体")
    elif level == 3:
        # Subsection heading: SimHei 12pt (小四号), left-aligned
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)
        run = para.add_run(text)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)  # 小四号
        run.font.color.rgb = RGBColor(0, 0, 0)
        set_east_asian_font(run, "黑体")
    else:
        # Sub-subsection: SimHei 12pt, left-aligned, not bold
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(3)
        run = para.add_run(text)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
        set_east_asian_font(run, "黑体")

    return para


def add_body_paragraph(doc, text):
    """Add a body paragraph with first-line indent."""
    para = doc.add_paragraph(text)
    para.style = doc.styles["Normal"]
    return para


def add_reference_paragraph(doc, text):
    """Add a reference entry with hanging indent, smaller font."""
    para = doc.add_paragraph(text)
    para.style = doc.styles["Normal"]
    pf = para.paragraph_format
    pf.first_line_indent = Cm(-1.27)  # Hanging indent
    pf.left_indent = Cm(1.27)
    pf.line_spacing = 1.5
    # Smaller font for references
    for run in para.runs:
        run.font.size = Pt(10.5)  # 五号
    return para


def parse_markdown(md_text):
    """Parse markdown text into structured sections."""
    sections = []
    current_section = {"type": "body", "heading": None, "level": 0, "paragraphs": []}
    buffer = []

    for line in md_text.split("\n"):
        stripped = line.strip()

        # Detect headings
        heading_match = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading_match:
            # Flush buffer
            if buffer:
                text = " ".join(buffer).strip()
                if text:
                    current_section["paragraphs"].append(text)
                buffer = []
            # Save current section
            if current_section["heading"] or current_section["paragraphs"]:
                sections.append(current_section)
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()
            # Detect reference section
            is_ref = any(
                kw in heading_text.lower()
                for kw in ["reference", "参考文献", "bibliography"]
            )
            current_section = {
                "type": "references" if is_ref else "body",
                "heading": heading_text,
                "level": level,
                "paragraphs": [],
            }
            continue

        # Blank line = paragraph break
        if not stripped:
            if buffer:
                text = " ".join(buffer).strip()
                if text:
                    current_section["paragraphs"].append(text)
                buffer = []
            continue

        buffer.append(stripped)

    # Flush remaining
    if buffer:
        text = " ".join(buffer).strip()
        if text:
            current_section["paragraphs"].append(text)
    if current_section["heading"] or current_section["paragraphs"]:
        sections.append(current_section)

    return sections


def generate_docx(md_text, output_path, title=None, chapter_num=2):
    """Generate a Word document from markdown text."""
    doc = Document()

    # Page setup (PhD dissertation standard)
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    set_normal_style(doc)

    # Chapter title
    if title:
        chapter_title = f"第{_num_to_chinese(chapter_num)}章 {title}"
    else:
        chapter_title = f"第{_num_to_chinese(chapter_num)}章 文献综述"
    add_chapter_title(doc, chapter_title)

    sections = parse_markdown(md_text)

    for sec in sections:
        # Add heading
        if sec["heading"]:
            level = sec["level"]
            if level == 1:
                # Skip top-level heading (already handled as chapter title)
                continue
            add_section_heading(doc, sec["heading"], level=level)

        # Add paragraphs
        if sec["type"] == "references":
            for para_text in sec["paragraphs"]:
                add_reference_paragraph(doc, para_text)
        else:
            for para_text in sec["paragraphs"]:
                add_body_paragraph(doc, para_text)

    doc.save(output_path)
    return output_path


def _num_to_chinese(n):
    """Convert number to Chinese character for chapter numbering."""
    mapping = {
        1: "一", 2: "二", 3: "三", 4: "四", 5: "五",
        6: "六", 7: "七", 8: "八", 9: "九", 10: "十",
    }
    return mapping.get(n, str(n))


def main():
    parser = argparse.ArgumentParser(
        description="Generate Word document for PhD dissertation literature review"
    )
    parser.add_argument("--input", required=True, help="Input markdown file path")
    parser.add_argument("--output", required=True, help="Output .docx file path")
    parser.add_argument("--title", default=None, help="Chapter title (e.g., '文献综述')")
    parser.add_argument(
        "--chapter-num",
        type=int,
        default=2,
        help="Chapter number (default: 2)",
    )
    args = parser.parse_args()

    try:
        with open(args.input, "r", encoding="utf-8") as f:
            md_text = f.read()
    except FileNotFoundError:
        print(f"Error: Input file '{args.input}' not found.", file=sys.stderr)
        sys.exit(1)

    output = generate_docx(md_text, args.output, title=args.title, chapter_num=args.chapter_num)
    print(f"Word document generated: {output}")


if __name__ == "__main__":
    main()
