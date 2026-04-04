"""
ta_agent.py — TA Research Agent 主引擎

将12个检查点的全自动流程封装为真正的 Python agent。
每个检查点独立调用 Anthropic API，有自己的 tool_use 循环。
支持断点续跑、自动重试、降级处理。

用法：
    python3 ta_agent.py --project-dir /path/to/project [--resume] [--from-cp N]
"""

import argparse
import os
import re
import sys
from pathlib import Path

import anthropic

# 确保同目录下的模块可以 import
sys.path.insert(0, str(Path(__file__).parent))

from ta_state import StateManager
from ta_tools import TOOL_DEFINITIONS, ToolExecutor, format_tool_result
from ta_checkpoints import CHECKPOINTS, build_system_prompt, build_task_message


# ─────────────────────────────────────────────
# 自动决策逻辑（Python 代码级，不依赖 Claude）
# ─────────────────────────────────────────────

# CP2：理论定位 A/B/C 关键词匹配
_THEORY_A_KEYWORDS = [
    "场域", "惯习", "资本", "结构化", "实践论", "行动者网络", "符号互动",
    "制度逻辑", "新制度", "布迪厄", "吉登斯", "福柯", "卢曼", "哈贝马斯",
    "bourdieu", "giddens", "foucault", "luhmann", "habermas", "ANT",
    "platform theory", "surveillance capitalism", "assemblage"
]
_MECHANISM_KEYWORDS = ["如何", "为什么", "机制", "过程", "路径", "影响", "作用", "how", "why", "mechanism"]
_DESCRIPTIVE_KEYWORDS = ["是什么", "有哪些", "现状", "特征", "类型", "差异", "what", "which", "describe"]


def cp2_theory_decision(research_info: dict) -> str:
    """
    根据研究问题自动判断理论定位 A/B/C。
    返回 "A"、"B" 或 "C"。
    """
    text = (research_info.get("topic", "") + " " + research_info.get("questions", "")).lower()

    # 优先级1：明确提及特定理论 → A
    if any(kw.lower() in text for kw in _THEORY_A_KEYWORDS):
        return "A"

    # 优先级2：机制性研究问题 → C
    if any(kw in text for kw in _MECHANISM_KEYWORDS):
        return "C"

    # 优先级3：描述性研究问题 → B
    if any(kw in text for kw in _DESCRIPTIVE_KEYWORDS):
        return "B"

    # 默认 C（最灵活，允许涌现）
    return "C"


def cp10_journal_decision(target_journal: str) -> str:
    """根据目标期刊类型确定标题/摘要/关键词方案"""
    if "ssci" in target_journal.lower():
        return "SSCI"
    return "C刊"


# ─────────────────────────────────────────────
# 检查点运行器（单个 CP 的 agent 调用）
# ─────────────────────────────────────────────

class CheckpointRunner:
    def __init__(self, cp_num: int, state: StateManager, client: anthropic.Anthropic):
        self.cp_num = cp_num
        self.cp_config = CHECKPOINTS[cp_num]
        self.state = state
        self.client = client
        self.tool_executor = ToolExecutor(state.project_dir)
        self.max_turns = self.cp_config.get("max_turns", 20)

    @property
    def system_prompt(self) -> str:
        return build_system_prompt(self.cp_num, self.cp_config.get("calls_skills", []))

    def build_task_message(self) -> str:
        return build_task_message(
            self.cp_num,
            self.state.get_research_info(),
            self.state.get_context(self.cp_num),
            self.state.get_auto_decisions()
        )

    def run(self) -> dict:
        """执行 tool_use 循环直到完成"""
        print(f"  → 启动子 agent（最多 {self.max_turns} 轮）...")

        messages = [{"role": "user", "content": self.build_task_message()}]
        turn = 0

        while turn < self.max_turns:
            resp = self.client.messages.create(
                model=os.environ.get("TA_AGENT_MODEL", "claude-opus-4-6"),
                max_tokens=8192,
                system=self.system_prompt,
                tools=TOOL_DEFINITIONS,
                messages=messages,
            )

            turn += 1

            # 终止条件1：end_turn
            if resp.stop_reason == "end_turn":
                print(f"  → CP{self.cp_num} 完成（{turn} 轮）")
                break

            # 提取 tool_use blocks
            tool_uses = [b for b in resp.content if b.type == "tool_use"]

            # 终止条件2：没有 tool_use（Claude 已完成所有操作）
            if not tool_uses:
                print(f"  → CP{self.cp_num} 完成（{turn} 轮，无更多工具调用）")
                break

            # 执行工具调用
            tool_results = []
            for tu in tool_uses:
                result = self.tool_executor.execute(tu.name, tu.input)
                tool_results.append(format_tool_result(tu.id, result))
                print(f"  ✓ {tu.name}({list(tu.input.keys())[0] if tu.input else ''})")

            # 继续对话
            messages.extend([
                {"role": "assistant", "content": resp.content},
                {"role": "user", "content": tool_results}
            ])

        # 终止条件3：超出 max_turns（降级处理，由调用方决定）
        if turn >= self.max_turns:
            print(f"  ⚠️  CP{self.cp_num} 达到最大轮次限制（{self.max_turns}）")

        # 收集已写入的输出文件
        outputs = self._collect_outputs()
        summary = self._extract_summary(messages)
        return {"outputs": outputs, "summary": summary}

    def _collect_outputs(self) -> list:
        """收集预期输出文件中已存在的文件路径"""
        expected = self.cp_config.get("expected_outputs", [])
        found = []
        for fname in expected:
            fpath = self.state.project_dir / fname
            if fpath.exists():
                found.append(fname)
        return found

    def _extract_summary(self, messages: list) -> str:
        """从最后一条 assistant 消息中提取一句话摘要"""
        for msg in reversed(messages):
            if msg.get("role") == "assistant":
                content = msg.get("content", "")
                if isinstance(content, list):
                    text = " ".join(
                        b.text for b in content if hasattr(b, "text")
                    )
                else:
                    text = str(content)
                # 取第一句话作为摘要（不超过100字）
                first_line = text.strip().split("\n")[0]
                return first_line[:100]
        return ""


# ─────────────────────────────────────────────
# 主协调器
# ─────────────────────────────────────────────

class TAAgent:
    def __init__(self, project_dir: str):
        self.state = StateManager(project_dir)
        self.client = self._init_client()

    def _init_client(self) -> anthropic.Anthropic:
        """初始化 Anthropic 客户端，支持自定义 base_url"""
        api_key = (
            os.environ.get("ANTHROPIC_API_KEY")
            or os.environ.get("ANTHROPIC_AUTH_TOKEN")
        )
        base_url = os.environ.get("ANTHROPIC_BASE_URL")

        kwargs = {}
        if api_key:
            kwargs["api_key"] = api_key
        if base_url:
            kwargs["base_url"] = base_url

        return anthropic.Anthropic(**kwargs)

    def initialize(self):
        """从 CLAUDE.md 读取研究信息（若尚未初始化）"""
        ri = self.state.get_research_info()
        if ri.get("topic"):
            return  # 已有研究信息，跳过

        claude_md = self.state.project_dir / "CLAUDE.md"
        if claude_md.exists():
            content = claude_md.read_text(encoding="utf-8")
            topic, questions, journal, interview_files = self._parse_claude_md(content)
        else:
            topic = input("请输入研究主题：").strip() or "未知研究主题"
            questions = input("请输入研究问题：").strip() or "未知研究问题"
            journal = input("目标期刊类型（C刊/SSCI，默认C刊）：").strip() or "C刊"
            interview_files = []

        self.state.set_research_info(topic, questions, journal, interview_files)
        print(f"\n✅ 研究信息已初始化：")
        print(f"   主题：{topic}")
        print(f"   期刊：{journal}")

    def _parse_claude_md(self, content: str) -> tuple:
        """从 CLAUDE.md 中提取研究基本信息（简单关键词匹配）"""
        topic = ""
        questions = ""
        journal = "C刊"
        interview_files = []

        lines = content.split("\n")
        for i, line in enumerate(lines):
            line_l = line.lower()
            if any(kw in line_l for kw in ["研究主题", "研究题目", "论文题目"]):
                topic = self._extract_value(line, lines, i)
            elif any(kw in line_l for kw in ["研究问题", "核心问题", "研究问"]):
                questions = self._extract_value(line, lines, i)
            elif any(kw in line_l for kw in ["目标期刊", "期刊", "发表目标"]):
                val = self._extract_value(line, lines, i).lower()
                if "ssci" in val:
                    journal = "SSCI"
            elif any(kw in line_l for kw in ["访谈", "interview", "材料路径"]):
                # 提取文件路径（简单模式匹配）
                paths = re.findall(r"[/~][\w/._-]+\.(?:txt|md|docx)", line)
                interview_files.extend(paths)

        # 若未解析到，使用文件前几行内容作为主题
        if not topic:
            topic = content.split("\n")[0].lstrip("#").strip()[:50] or "未知主题"

        return topic, questions, journal, interview_files

    def _extract_value(self, line: str, lines: list, idx: int) -> str:
        """从行内或下一行提取值"""
        # 尝试从当前行提取（格式：xxx：value 或 xxx: value）
        for sep in ["：", ":", "="]:
            if sep in line:
                val = line.split(sep, 1)[1].strip().strip("*").strip()
                if val and len(val) > 1:
                    return val
        # 尝试下一行
        if idx + 1 < len(lines):
            next_line = lines[idx + 1].strip().lstrip("-").lstrip("*").strip()
            if next_line and not next_line.startswith("#"):
                return next_line[:200]
        return ""

    def pre_decision(self, cp_num: int):
        """执行 CP2/CP8/CP10 的 Python 级自动决策"""
        decisions = self.state.get_auto_decisions()

        if cp_num == 2 and not decisions.get("cp2_theory"):
            theory = cp2_theory_decision(self.state.get_research_info())
            self.state.record_decision("cp2_theory", theory)
            label = {"A": "理论驱动", "B": "经验驱动", "C": "敏感性概念"}[theory]
            print(f"  🤖 自动决策（CP2）：理论定位 = {theory}（{label}）")

        elif cp_num == 8 and not decisions.get("cp8_narrative"):
            # CP8 叙事结构：让子 agent 分析后在输出文件中标注，Python 读取并记录
            # 此处先设默认值，子 agent 执行完后从输出中提取
            self.state.record_decision("cp8_narrative", "parallel")  # 先设默认

        elif cp_num == 10 and not decisions.get("cp10_journal_type"):
            journal_type = cp10_journal_decision(self.state.get_target_journal())
            self.state.record_decision("cp10_journal_type", journal_type)
            print(f"  🤖 自动决策（CP10）：期刊类型 = {journal_type}")

    def post_decision(self, cp_num: int):
        """CP 完成后更新决策（如从输出文件中提取 CP8 叙事结构）"""
        if cp_num == 8:
            findings_path = self.state.project_dir / "findings.md"
            if findings_path.exists():
                content = findings_path.read_text(encoding="utf-8")
                # 提取标注：[叙事结构] 并列型/串联型 | 依据：XXX
                match = re.search(r"\[叙事结构\]\s*(并列型|串联型)", content)
                if match:
                    narrative = "parallel" if "并列" in match.group(1) else "serial"
                    self.state.record_decision("cp8_narrative", narrative)
                    print(f"  🤖 更新决策（CP8）：叙事结构 = {narrative}")

    def run(self, from_cp: int = 1):
        """主循环：按顺序执行 CP1-CP12"""
        print(f"\n{'='*60}")
        print(f"[TA-Agent] 全自动 TA 研究工作流启动")
        print(f"项目目录：{self.state.project_dir}")
        print(f"{'='*60}")

        self.initialize()
        self.state.print_status()

        for cp_num in range(from_cp, 13):
            if self.state.is_completed(cp_num):
                print(f"  ↳ CP{cp_num} 已完成，跳过")
                continue

            cp_name = CHECKPOINTS[cp_num]["name"]
            print(f"\n{'─'*60}")
            print(f"▶ 检查点 {cp_num}/12：{cp_name}")
            print(f"{'─'*60}")

            self.pre_decision(cp_num)
            result = self.run_with_retry(cp_num)
            self.post_decision(cp_num)

            self.state.mark_complete(
                cp_num,
                result.get("outputs", []),
                result.get("summary", "")
            )

            print(f"✅ CP{cp_num} 完成 → 输出文件：{result.get('outputs', [])}")

        print(f"\n{'─'*60}")
        print(f"▶ 后处理：合并章节 → Word 初稿")
        print(f"{'─'*60}")
        word_path = self.assemble_word_draft()

        print(f"\n{'='*60}")
        print(f"[TA-Agent] 全部12个检查点已完成 ✅")
        print(f"最终报告：{self.state.project_dir}/final_report.md")
        if word_path:
            print(f"Word 初稿：{word_path}")
        print(f"{'='*60}")
        self.state.print_status()

    def assemble_word_draft(self) -> str:
        """将各章节 md 文件按论文顺序合并为完整 Word 初稿（.docx）"""
        chapter_files = [
            "title_abstract_keywords.md",
            "introduction.md",
            "literature_review.md",
            "framework_section.md",
            "methods.md",
            "findings.md",
            "discussion.md",
        ]

        available = [
            self.state.project_dir / f
            for f in chapter_files
            if (self.state.project_dir / f).exists()
        ]

        if not available:
            print("  ⚠️  未找到任何章节文件，跳过 Word 生成")
            return ""

        ri = self.state.get_research_info()
        topic = ri.get("topic", "论文初稿")
        safe_topic = re.sub(r'[/\\:*?"<>|]', "_", topic)

        # ── 方案1：pypandoc（最佳格式质量）──────────────────────
        try:
            import pypandoc
            combined_md = self.state.project_dir / "_draft_combined.md"
            with open(combined_md, "w", encoding="utf-8") as out:
                for fpath in available:
                    out.write(fpath.read_text(encoding="utf-8"))
                    out.write("\n\n\n")
            output_path = self.state.project_dir / f"{safe_topic}_初稿.docx"
            pypandoc.convert_file(str(combined_md), "docx", outputfile=str(output_path))
            combined_md.unlink()
            print(f"  ✅ Word 初稿已生成（pypandoc）：{output_path.name}")
            return str(output_path)
        except ImportError:
            pass
        except Exception as e:
            print(f"  ⚠️  pypandoc 转换失败：{e}，尝试备选方案")

        # ── 方案2：python-docx（基础 markdown 解析）─────────────
        try:
            from docx import Document
            from docx.shared import Pt

            doc = Document()

            # 设置正文字体
            normal = doc.styles["Normal"]
            normal.font.name = "宋体"
            normal.font.size = Pt(12)

            for fpath in available:
                content = fpath.read_text(encoding="utf-8")
                self._md_to_docx(doc, content)
                doc.add_page_break()

            output_path = self.state.project_dir / f"{safe_topic}_初稿.docx"
            doc.save(str(output_path))
            print(f"  ✅ Word 初稿已生成（python-docx）：{output_path.name}")
            return str(output_path)
        except ImportError:
            pass

        # ── 方案3：降级为合并 md ─────────────────────────────────
        output_path = self.state.project_dir / f"{safe_topic}_初稿.md"
        with open(output_path, "w", encoding="utf-8") as out:
            for fpath in available:
                out.write(fpath.read_text(encoding="utf-8"))
                out.write("\n\n---\n\n")
        print(f"  ⚠️  未安装 python-docx/pypandoc，已生成合并 md：{output_path.name}")
        print(f"       安装 Word 支持：pip install python-docx")
        return str(output_path)

    def _md_to_docx(self, doc, markdown_text: str):
        """将 markdown 文本转为 docx 段落（基础实现）"""
        from docx.shared import Pt

        for line in markdown_text.split("\n"):
            if line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:].strip(), level=1)
            elif line.startswith(("- ", "* ")):
                doc.add_paragraph(line[2:].strip(), style="List Bullet")
            elif line.strip() in ("", "---"):
                pass
            else:
                # 去除行内 markdown 标记，保留纯文本
                clean = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
                clean = re.sub(r"\*(.+?)\*", r"\1", clean)
                clean = re.sub(r"`(.+?)`", r"\1", clean)
                clean = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", clean)
                if clean.strip():
                    doc.add_paragraph(clean.strip())

    def run_with_retry(self, cp_num: int, max_attempts: int = 3) -> dict:
        """带重试的检查点执行"""
        for attempt in range(1, max_attempts + 1):
            try:
                runner = CheckpointRunner(cp_num, self.state, self.client)
                return runner.run()
            except anthropic.APIError as e:
                print(f"  ⚠️  API 错误（尝试 {attempt}/{max_attempts}）：{e}")
                if attempt == max_attempts:
                    return self.degrade(cp_num, str(e), attempt)
            except Exception as e:
                print(f"  ⚠️  执行错误（尝试 {attempt}/{max_attempts}）：{e}")
                if attempt == max_attempts:
                    return self.degrade(cp_num, str(e), attempt)

        return {}

    def degrade(self, cp_num: int, error: str, attempts: int) -> dict:
        """降级处理：记录失败，生成说明文件，继续流程"""
        cp_name = CHECKPOINTS[cp_num]["name"]
        print(f"  ❌ CP{cp_num}（{cp_name}）降级处理：{error}")

        # 写一个说明文件
        fallback_file = self.state.project_dir / f"cp{cp_num}_skipped.md"
        fallback_file.write_text(
            f"# CP{cp_num} — {cp_name}（跳过）\n\n"
            f"**跳过原因**：{error}\n\n"
            f"**尝试次数**：{attempts}\n\n"
            f"**建议**：请手动执行此检查点，完成后将结果保存到对应文件。",
            encoding="utf-8"
        )
        self.state.mark_failed(cp_num, error, attempts)
        return {"outputs": [f"cp{cp_num}_skipped.md"], "summary": f"跳过：{error[:50]}"}


# ─────────────────────────────────────────────
# CLI 入口
# ─────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="TA Research Agent — 全自动学术论文写作工作流"
    )
    parser.add_argument(
        "--project-dir",
        required=True,
        help="项目目录路径（论文工作目录）"
    )
    parser.add_argument(
        "--resume",
        action="store_true",
        help="从上次中断的检查点继续"
    )
    parser.add_argument(
        "--from-cp",
        type=int,
        default=1,
        metavar="N",
        help="从第 N 个检查点开始（1-12）"
    )
    parser.add_argument(
        "--status",
        action="store_true",
        help="只显示当前进度，不执行"
    )

    args = parser.parse_args()

    # 验证项目目录
    project_dir = Path(args.project_dir).resolve()
    if not project_dir.exists():
        print(f"❌ 项目目录不存在：{project_dir}")
        sys.exit(1)

    agent = TAAgent(str(project_dir))

    if args.status:
        agent.state.print_status()
        return

    # 确定起始检查点
    if args.resume:
        completed = sorted(agent.state.state.get("completed_cps", []))
        from_cp = max(completed) + 1 if completed else 1
        print(f"[续跑模式] 从 CP{from_cp} 继续")
    else:
        from_cp = max(1, min(12, args.from_cp))

    agent.run(from_cp=from_cp)


if __name__ == "__main__":
    main()
