#!/usr/bin/env python3
"""
论文评阅修订工具：读取 Word 文档，根据 JSON 格式的评审意见，
生成带修订痕迹（tracked changes）和批注（comments）的新 Word 文档。
"""

import json
import sys
import copy
import re
from datetime import datetime
from pathlib import Path
from lxml import etree
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# Word Open XML 命名空间
NAMESPACES = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
}

AUTHOR = "学术评阅助手"
DATE_STR = datetime.now().strftime("%Y-%m-%dT%H:%M:00Z")


def make_w(tag):
    """创建带 w 命名空间的元素"""
    return etree.SubElement.__func__  # placeholder


def create_element(tag, nsmap=None):
    """创建 OOXML 元素"""
    ns = NAMESPACES['w']
    return etree.Element(f'{{{ns}}}{tag}', nsmap=nsmap)


def create_sub_element(parent, tag):
    """创建子元素"""
    ns = NAMESPACES['w']
    return etree.SubElement(parent, f'{{{ns}}}{tag}')


def qn(tag):
    """生成带命名空间的完整标签名"""
    prefix, local = tag.split(':')
    return f'{{{NAMESPACES[prefix]}}}{local}'


def get_revision_id():
    """生成递增的修订 ID"""
    if not hasattr(get_revision_id, '_counter'):
        get_revision_id._counter = 100
    get_revision_id._counter += 1
    return str(get_revision_id._counter)


def add_comment_to_document(document, comment_text, author=AUTHOR, initials="审"):
    """
    向文档的 comments part 添加一条评论，返回 comment_id
    """
    # 获取或创建 comments part
    comments_part = None
    for rel in document.part.rels.values():
        if "comments" in rel.reltype:
            comments_part = rel.target_part
            break

    if comments_part is None:
        comments_part = _create_comments_part(document)

    comments_xml = etree.fromstring(comments_part.blob)
    # 计算新 comment id
    existing_ids = comments_xml.findall(qn('w:comment'))
    comment_id = str(len(existing_ids))

    # 创建 comment 元素
    comment_el = etree.SubElement(comments_xml, qn('w:comment'))
    comment_el.set(qn('w:id'), comment_id)
    comment_el.set(qn('w:author'), author)
    comment_el.set(qn('w:date'), DATE_STR)
    comment_el.set(qn('w:initials'), initials)

    # comment 内容
    p = etree.SubElement(comment_el, qn('w:p'))
    r = etree.SubElement(p, qn('w:r'))
    t = etree.SubElement(r, qn('w:t'))
    t.text = comment_text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    comments_part._blob = etree.tostring(comments_xml, xml_declaration=True,
                                          encoding='UTF-8', standalone=True)
    return comment_id


def _create_comments_part(document):
    """创建 comments.xml part"""
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI

    nsmap = {
        'w': NAMESPACES['w'],
        'r': NAMESPACES['r'],
    }
    comments_xml = etree.Element(qn('w:comments'), nsmap=nsmap)

    blob = etree.tostring(comments_xml, xml_declaration=True,
                          encoding='UTF-8', standalone=True)

    part_name = PackURI('/word/comments.xml')
    content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml'
    comments_part = Part(part_name, content_type, blob, document.part.package)

    document.part.relate_to(comments_part,
                            'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments')
    return comments_part


def wrap_paragraph_with_comment(paragraph, comment_id):
    """在段落上添加批注引用标记"""
    p_element = paragraph._element

    # commentRangeStart 放在段首
    range_start = etree.Element(qn('w:commentRangeStart'))
    range_start.set(qn('w:id'), comment_id)
    p_element.insert(0, range_start)

    # commentRangeEnd 放在段尾
    range_end = etree.SubElement(p_element, qn('w:commentRangeEnd'))
    range_end.set(qn('w:id'), comment_id)

    # commentReference run
    r = etree.SubElement(p_element, qn('w:r'))
    rpr = etree.SubElement(r, qn('w:rPr'))
    rstyle = etree.SubElement(rpr, qn('w:rStyle'))
    rstyle.set(qn('w:val'), 'CommentReference')
    ref = etree.SubElement(r, qn('w:commentReference'))
    ref.set(qn('w:id'), comment_id)


def apply_text_replacement(paragraph, old_text, new_text):
    """
    在段落中用修订痕迹（tracked changes）替换文本。
    将匹配的文本标记为删除（红色删除线），新文本标记为插入（红色下划线）。
    """
    p_element = paragraph._element

    # 收集所有 run 的文本和对应元素
    runs = p_element.findall(qn('w:r'))
    if not runs:
        return False

    # 拼接完整段落文本
    full_text = ""
    run_map = []  # [(run_element, text_element, start_pos, end_pos)]
    for run in runs:
        t_elements = run.findall(qn('w:t'))
        for t in t_elements:
            text = t.text or ""
            start = len(full_text)
            full_text += text
            run_map.append((run, t, start, start + len(text)))

    # 查找 old_text 位置
    match_start = full_text.find(old_text)
    if match_start == -1:
        # 尝试忽略空白匹配
        normalized_full = re.sub(r'\s+', '', full_text)
        normalized_old = re.sub(r'\s+', '', old_text)
        norm_pos = normalized_full.find(normalized_old)
        if norm_pos == -1:
            return False
        # 映射回原始位置（近似）
        match_start = full_text.find(old_text[:10])
        if match_start == -1:
            match_start = 0

    match_end = match_start + len(old_text)
    rev_id_del = get_revision_id()
    rev_id_ins = get_revision_id()

    # 创建删除标记
    del_start = etree.Element(qn('w:del'))
    del_start.set(qn('w:id'), rev_id_del)
    del_start.set(qn('w:author'), AUTHOR)
    del_start.set(qn('w:date'), DATE_STR)

    del_run = etree.SubElement(del_start, qn('w:r'))
    # 复制原始 run 的格式
    for run, t, s, e in run_map:
        if s < match_end and e > match_start:
            rpr = run.find(qn('w:rPr'))
            if rpr is not None:
                del_run.append(copy.deepcopy(rpr))
            break

    del_text = etree.SubElement(del_run, qn('w:delText'))
    del_text.text = old_text
    del_text.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    # 创建插入标记
    ins_start = etree.Element(qn('w:ins'))
    ins_start.set(qn('w:id'), rev_id_ins)
    ins_start.set(qn('w:author'), AUTHOR)
    ins_start.set(qn('w:date'), DATE_STR)

    ins_run = etree.SubElement(ins_start, qn('w:r'))
    # 复制格式
    for run, t, s, e in run_map:
        if s < match_end and e > match_start:
            rpr = run.find(qn('w:rPr'))
            if rpr is not None:
                ins_run.append(copy.deepcopy(rpr))
            break

    ins_text = etree.SubElement(ins_run, qn('w:t'))
    ins_text.text = new_text
    ins_text.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    # 找到需要替换的 run 位置并插入修订标记
    insert_point = None
    runs_to_handle = []
    for run, t, s, e in run_map:
        if s < match_end and e > match_start:
            runs_to_handle.append((run, t, s, e))
            if insert_point is None:
                insert_point = run

    if insert_point is None:
        return False

    # 在第一个匹配 run 前插入 del 和 ins
    parent = insert_point.getparent()
    idx = list(parent).index(insert_point)

    # 处理：如果 old_text 不覆盖整个 run，需要拆分
    first_run, first_t, first_s, first_e = runs_to_handle[0]

    if first_s < match_start:
        # run 开头有不需要删除的文本
        prefix_text = full_text[first_s:match_start]
        first_t.text = prefix_text

        parent.insert(idx + 1, del_start)
        parent.insert(idx + 2, ins_start)

        # 处理 run 结尾的剩余文本
        last_run, last_t, last_s, last_e = runs_to_handle[-1]
        if last_e > match_end:
            suffix_text = full_text[match_end:last_e]
            suffix_run = copy.deepcopy(last_run)
            for st in suffix_run.findall(qn('w:t')):
                st.text = suffix_text
            tail_idx = list(parent).index(ins_start) + 1
            parent.insert(tail_idx, suffix_run)

        # 删除中间被完全覆盖的 run
        for run, t, s, e in runs_to_handle[1:]:
            if run.getparent() is not None:
                parent.remove(run)
    else:
        parent.insert(idx, del_start)
        parent.insert(idx + 1, ins_start)

        # 删除被替换的 run
        for run, t, s, e in runs_to_handle:
            if run.getparent() is not None:
                parent.remove(run)

        # 如果最后一个 run 有剩余文本
        last_run, last_t, last_s, last_e = runs_to_handle[-1]
        if last_e > match_end:
            suffix_text = full_text[match_end:last_e]
            suffix_run = copy.deepcopy(last_run)
            for st in suffix_run.findall(qn('w:t')):
                st.text = suffix_text
            tail_idx = list(parent).index(ins_start) + 1
            parent.insert(tail_idx, suffix_run)

    return True


def _enable_revision_display(document):
    """
    修改 settings.xml，确保 Word 打开文档时默认显示修订痕迹和批注。
    添加 <w:trackRevisions/> 和 <w:revisionView> 元素。
    """
    settings_element = document.settings.element
    ns = NAMESPACES['w']

    # 添加 <w:trackRevisions/> — 告知 Word 文档启用了修订跟踪
    if settings_element.find(qn('w:trackRevisions')) is None:
        track_rev = etree.SubElement(settings_element, qn('w:trackRevisions'))

    # 移除可能存在的 <w:revisionView> 元素（某些模板可能隐藏标记）
    for rv in settings_element.findall(qn('w:revisionView')):
        settings_element.remove(rv)

    # 添加 <w:revisionView> 显示所有修订类型
    rev_view = etree.SubElement(settings_element, qn('w:revisionView'))
    rev_view.set(qn('w:markup'), '1')
    rev_view.set(qn('w:comments'), '1')
    rev_view.set(qn('w:formatting'), '1')
    rev_view.set(qn('w:inkAnnotations'), '1')
    rev_view.set(qn('w:insDel'), '1')


###############################################################################
# 可视化格式标注模式 — 直接用颜色和格式标注修订，不依赖 Word 修订引擎
###############################################################################

def _visual_replace(paragraph, old_text, new_text):
    """
    在段落中找到 old_text，将其标记为红色删除线，在其后插入蓝色下划线的 new_text。
    返回 True/False 表示是否成功。
    """
    p_element = paragraph._element

    runs = p_element.findall(qn('w:r'))
    if not runs:
        return False

    full_text = ""
    run_map = []
    for run in runs:
        for t in run.findall(qn('w:t')):
            text = t.text or ""
            start = len(full_text)
            full_text += text
            run_map.append((run, t, start, start + len(text)))

    match_start = full_text.find(old_text)
    if match_start == -1:
        return False
    match_end = match_start + len(old_text)

    affected = [(run, t, s, e) for run, t, s, e in run_map
                if s < match_end and e > match_start]
    if not affected:
        return False

    first_run, first_t, first_s, first_e = affected[0]
    last_run, last_t, last_s, last_e = affected[-1]

    orig_rpr = first_run.find(qn('w:rPr'))

    def _make_rpr():
        return copy.deepcopy(orig_rpr) if orig_rpr is not None else etree.Element(qn('w:rPr'))

    # 红色删除线 run
    del_run = etree.Element(qn('w:r'))
    del_rpr = _make_rpr()
    etree.SubElement(del_rpr, qn('w:strike')).set(qn('w:val'), 'true')
    c = del_rpr.find(qn('w:color'))
    if c is None:
        c = etree.SubElement(del_rpr, qn('w:color'))
    c.set(qn('w:val'), 'FF0000')
    del_run.insert(0, del_rpr)
    dt = etree.SubElement(del_run, qn('w:t'))
    dt.text = old_text
    dt.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    # 蓝色下划线 run
    ins_run = etree.Element(qn('w:r'))
    ins_rpr = _make_rpr()
    u = ins_rpr.find(qn('w:u'))
    if u is None:
        u = etree.SubElement(ins_rpr, qn('w:u'))
    u.set(qn('w:val'), 'single')
    u.set(qn('w:color'), '0000FF')
    c2 = ins_rpr.find(qn('w:color'))
    if c2 is None:
        c2 = etree.SubElement(ins_rpr, qn('w:color'))
    c2.set(qn('w:val'), '0000FF')
    ins_run.insert(0, ins_rpr)
    it = etree.SubElement(ins_run, qn('w:t'))
    it.text = new_text
    it.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    # 插入到段落
    parent = first_run.getparent()

    if first_s < match_start:
        first_t.text = full_text[first_s:match_start]
        idx = list(parent).index(first_run) + 1
        parent.insert(idx, del_run)
        parent.insert(idx + 1, ins_run)
        if last_e > match_end:
            suffix_run = copy.deepcopy(last_run)
            for st in suffix_run.findall(qn('w:t')):
                st.text = full_text[match_end:last_e]
            parent.insert(idx + 2, suffix_run)
        for run, t, s, e in affected[1:]:
            if run.getparent() is not None:
                parent.remove(run)
    else:
        idx = list(parent).index(first_run)
        parent.insert(idx, del_run)
        parent.insert(idx + 1, ins_run)
        for run, t, s, e in affected:
            if run.getparent() is not None:
                parent.remove(run)
        if last_e > match_end:
            suffix_run = copy.deepcopy(last_run)
            for st in suffix_run.findall(qn('w:t')):
                st.text = full_text[match_end:last_e]
            parent.insert(list(parent).index(ins_run) + 1, suffix_run)

    return True


def _visual_comment(paragraph, comment_text):
    """
    在段落后插入一个橙色边框 + 浅黄色底纹的审稿意见段落。
    """
    p_el = paragraph._element
    parent = p_el.getparent()
    idx = list(parent).index(p_el)

    new_p = etree.Element(qn('w:p'))

    # 段落格式
    pPr = etree.SubElement(new_p, qn('w:pPr'))
    pBdr = etree.SubElement(pPr, qn('w:pBdr'))
    left_bdr = etree.SubElement(pBdr, qn('w:left'))
    left_bdr.set(qn('w:val'), 'single')
    left_bdr.set(qn('w:sz'), '12')
    left_bdr.set(qn('w:space'), '4')
    left_bdr.set(qn('w:color'), 'FF8C00')
    shd = etree.SubElement(pPr, qn('w:shd'))
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFF8DC')
    rPr_default = etree.SubElement(pPr, qn('w:rPr'))
    etree.SubElement(rPr_default, qn('w:sz')).set(qn('w:val'), '18')
    etree.SubElement(rPr_default, qn('w:szCs')).set(qn('w:val'), '18')

    # 标签 run
    label_run = etree.SubElement(new_p, qn('w:r'))
    label_rpr = etree.SubElement(label_run, qn('w:rPr'))
    etree.SubElement(label_rpr, qn('w:b'))
    etree.SubElement(label_rpr, qn('w:color')).set(qn('w:val'), 'D2691E')
    etree.SubElement(label_rpr, qn('w:sz')).set(qn('w:val'), '18')
    etree.SubElement(label_rpr, qn('w:szCs')).set(qn('w:val'), '18')
    lt = etree.SubElement(label_run, qn('w:t'))
    lt.text = "【审稿意见】"
    lt.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    # 内容 run
    content_run = etree.SubElement(new_p, qn('w:r'))
    content_rpr = etree.SubElement(content_run, qn('w:rPr'))
    etree.SubElement(content_rpr, qn('w:color')).set(qn('w:val'), '333333')
    etree.SubElement(content_rpr, qn('w:sz')).set(qn('w:val'), '18')
    etree.SubElement(content_rpr, qn('w:szCs')).set(qn('w:val'), '18')
    ct = etree.SubElement(content_run, qn('w:t'))
    ct.text = comment_text
    ct.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    parent.insert(idx + 1, new_p)


def apply_visual(input_path, revisions_json, output_path=None):
    """
    可视化格式标注模式：用颜色和格式直接在文档中标注修订。
    - 删除文本：红色 + 删除线
    - 新增文本：蓝色 + 下划线
    - 审稿意见：橙色边框 + 浅黄底纹段落
    兼容所有文字处理软件（Word、WPS、LibreOffice、Pages）。
    """
    input_path = Path(input_path)
    if output_path is None:
        output_path = input_path.parent / f"{input_path.stem}_可视化评阅{input_path.suffix}"
    else:
        output_path = Path(output_path)

    doc = Document(str(input_path))

    if isinstance(revisions_json, str):
        revisions = json.loads(revisions_json)
    else:
        revisions = revisions_json

    paragraphs = doc.paragraphs
    total = len(revisions.get("revisions", []))
    applied = 0
    skipped = 0

    # 先处理 replace 类型（不改变段落数量）
    comment_insertions = []  # [(para_idx, comment_text)]

    for rev in revisions.get("revisions", []):
        rev_type = rev.get("type", "comment")
        para_idx = rev.get("paragraph_index", -1)

        if para_idx < 0 or para_idx >= len(paragraphs):
            old_text = rev.get("old_text", "")
            if old_text:
                for i, p in enumerate(paragraphs):
                    if old_text[:20] in p.text:
                        para_idx = i
                        break

        if para_idx < 0 or para_idx >= len(paragraphs):
            skipped += 1
            continue

        if rev_type == "replace":
            old_text = rev.get("old_text", "")
            new_text = rev.get("new_text", "")
            comment_text = rev.get("comment", "")

            if old_text and new_text:
                success = _visual_replace(paragraphs[para_idx], old_text, new_text)
                if success:
                    applied += 1
                    if comment_text:
                        comment_insertions.append((para_idx, comment_text))
                else:
                    if comment_text:
                        comment_insertions.append((para_idx,
                            f"[建议修改] {old_text} → {new_text}\n\n{comment_text}"))
                        applied += 1
                    else:
                        skipped += 1

        elif rev_type == "comment":
            comment_text = rev.get("comment", "")
            if comment_text:
                comment_insertions.append((para_idx, comment_text))
                applied += 1

        elif rev_type == "delete":
            old_text = rev.get("old_text", "")
            comment_text = rev.get("comment", "建议删除此段")
            comment_insertions.append((para_idx,
                f"[建议删除] {comment_text}"))
            # 将段落文字标记为红色删除线
            p_el = paragraphs[para_idx]._element
            for run in p_el.findall(qn('w:r')):
                rpr = run.find(qn('w:rPr'))
                if rpr is None:
                    rpr = etree.SubElement(run, qn('w:rPr'))
                    run.insert(0, rpr)
                etree.SubElement(rpr, qn('w:strike')).set(qn('w:val'), 'true')
                c = rpr.find(qn('w:color'))
                if c is None:
                    c = etree.SubElement(rpr, qn('w:color'))
                c.set(qn('w:val'), 'FF0000')
            applied += 1

    # 倒序插入审稿意见段落（避免索引偏移）
    comment_insertions.sort(key=lambda x: x[0], reverse=True)
    for para_idx, comment_text in comment_insertions:
        if para_idx < len(paragraphs):
            _visual_comment(paragraphs[para_idx], comment_text)

    doc.save(str(output_path))

    return {
        "output_path": str(output_path),
        "total_revisions": total,
        "applied": applied,
        "skipped": skipped
    }


###############################################################################
# 修订痕迹模式（tracked changes）— 依赖 Word 修订引擎
###############################################################################

def apply_revisions(input_path, revisions_json, output_path=None):
    """
    修订痕迹模式：生成带 OOXML tracked changes 和 comments 的文档。
    需要在 Word 中开启「审阅→显示标记」查看。WPS 兼容性有限。

    revisions_json 格式：
    {
        "revisions": [
            {
                "type": "replace",
                "paragraph_index": 3,
                "old_text": "原始文本片段",
                "new_text": "修改后的文本",
                "comment": "修改理由说明"
            },
            {
                "type": "comment",
                "paragraph_index": 5,
                "comment": "评阅意见"
            }
        ]
    }
    """
    input_path = Path(input_path)
    if output_path is None:
        output_path = input_path.parent / f"{input_path.stem}_评阅修订{input_path.suffix}"
    else:
        output_path = Path(output_path)

    doc = Document(str(input_path))

    if isinstance(revisions_json, str):
        revisions = json.loads(revisions_json)
    else:
        revisions = revisions_json

    paragraphs = doc.paragraphs
    total = len(revisions.get("revisions", []))
    applied = 0
    skipped = 0

    # 收集需要插入正文的评审意见段落（倒序插入避免索引偏移）
    comment_insertions = []  # [(para_idx, comment_text)]

    for rev in revisions.get("revisions", []):
        rev_type = rev.get("type", "comment")
        para_idx = rev.get("paragraph_index", -1)

        if para_idx < 0 or para_idx >= len(paragraphs):
            # 尝试通过文本内容查找段落
            old_text = rev.get("old_text", "")
            if old_text:
                for i, p in enumerate(paragraphs):
                    if old_text[:20] in p.text:
                        para_idx = i
                        break

        if para_idx < 0 or para_idx >= len(paragraphs):
            skipped += 1
            continue

        paragraph = paragraphs[para_idx]

        if rev_type == "replace":
            old_text = rev.get("old_text", "")
            new_text = rev.get("new_text", "")
            comment_text = rev.get("comment", "")

            if old_text and new_text:
                success = apply_text_replacement(paragraph, old_text, new_text)
                if success:
                    applied += 1
                    if comment_text:
                        comment_insertions.append((para_idx, comment_text))
                else:
                    # 替换失败，至少插入正文评语
                    if comment_text:
                        comment_insertions.append((para_idx,
                            f"[建议修改] {old_text} → {new_text}\n\n{comment_text}"))
                        applied += 1
                    else:
                        skipped += 1

        elif rev_type == "comment":
            comment_text = rev.get("comment", "")
            if comment_text:
                comment_insertions.append((para_idx, comment_text))
                applied += 1

        elif rev_type == "delete":
            old_text = rev.get("old_text", paragraph.text)
            comment_text = rev.get("comment", "建议删除此段")
            # 标记为删除
            rev_id = get_revision_id()
            p_el = paragraph._element
            del_el = etree.Element(qn('w:del'))
            del_el.set(qn('w:id'), rev_id)
            del_el.set(qn('w:author'), AUTHOR)
            del_el.set(qn('w:date'), DATE_STR)

            # 将段落中所有 run 移入 del 标记
            runs = list(p_el.findall(qn('w:r')))
            for run in runs:
                # 将 w:t 改为 w:delText
                for t in run.findall(qn('w:t')):
                    new_del_t = etree.SubElement(run, qn('w:delText'))
                    new_del_t.text = t.text
                    new_del_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                    run.remove(t)
                del_el.append(run)

            p_el.append(del_el)

            if comment_text:
                comment_insertions.append((para_idx, f"[建议删除] {comment_text}"))
            applied += 1

    # 倒序插入正文评审意见段落（复用可视化模式的格式），避免索引偏移
    comment_insertions.sort(key=lambda x: x[0], reverse=True)
    for para_idx, comment_text in comment_insertions:
        if para_idx < len(paragraphs):
            _visual_comment(paragraphs[para_idx], comment_text)

    # 在 settings.xml 中启用修订显示，确保 Word 打开时默认显示标记
    _enable_revision_display(doc)

    doc.save(str(output_path))

    return {
        "output_path": str(output_path),
        "total_revisions": total,
        "applied": applied,
        "skipped": skipped
    }


def convert_pdf_to_docx(pdf_path, docx_path=None):
    """将 PDF 转换为 DOCX 格式"""
    from pdf2docx import Converter

    pdf_path = Path(pdf_path)
    if docx_path is None:
        docx_path = pdf_path.parent / f"{pdf_path.stem}_转换.docx"
    else:
        docx_path = Path(docx_path)

    cv = Converter(str(pdf_path))
    cv.convert(str(docx_path))
    cv.close()

    return str(docx_path)


def ensure_docx(input_path):
    """
    确保输入为 docx 格式。如果是 PDF，自动转换。
    返回 (docx_path, was_converted)
    """
    input_path = Path(input_path)
    if input_path.suffix.lower() == '.pdf':
        docx_path = input_path.parent / f"{input_path.stem}_转换.docx"
        print(f"检测到 PDF 文件，正在转换为 DOCX: {docx_path}", file=sys.stderr)
        convert_pdf_to_docx(input_path, docx_path)
        print(f"转换完成: {docx_path}", file=sys.stderr)
        return str(docx_path), True
    return str(input_path), False


def extract_text(input_path):
    """提取文档的全部文本，带段落编号。支持 .docx 和 .pdf"""
    input_path = Path(input_path)

    if input_path.suffix.lower() == '.pdf':
        docx_path, _ = ensure_docx(input_path)
    else:
        docx_path = str(input_path)

    doc = Document(docx_path)
    result = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            style = para.style.name if para.style else "Normal"
            result.append({
                "index": i,
                "style": style,
                "text": text
            })
    return result


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法:")
        print("  提取文本: python apply_revisions.py extract <input.docx|pdf>")
        print("  可视化评阅（推荐）: python apply_revisions.py visual <input.docx|pdf> <revisions.json> [output.docx]")
        print("  修订痕迹: python apply_revisions.py apply <input.docx|pdf> <revisions.json> [output.docx]")
        print("  转换PDF:  python apply_revisions.py convert <input.pdf> [output.docx]")
        sys.exit(1)

    command = sys.argv[1]

    if command == "extract":
        input_file = sys.argv[2]
        paragraphs = extract_text(input_file)
        print(json.dumps(paragraphs, ensure_ascii=False, indent=2))

    elif command in ("visual", "apply"):
        input_file = sys.argv[2]
        revisions_file = sys.argv[3]
        output_file = sys.argv[4] if len(sys.argv) > 4 else None

        # 如果是 PDF，先转换
        input_path = Path(input_file)
        if input_path.suffix.lower() == '.pdf':
            docx_file, _ = ensure_docx(input_file)
            if output_file is None:
                suffix = "_可视化评阅.docx" if command == "visual" else "_评阅修订.docx"
                output_file = str(input_path.parent / f"{input_path.stem}{suffix}")
        else:
            docx_file = input_file

        with open(revisions_file, 'r', encoding='utf-8') as f:
            revisions = json.load(f)

        if command == "visual":
            result = apply_visual(docx_file, revisions, output_file)
        else:
            result = apply_revisions(docx_file, revisions, output_file)
        print(json.dumps(result, ensure_ascii=False, indent=2))

    elif command == "convert":
        input_file = sys.argv[2]
        output_file = sys.argv[3] if len(sys.argv) > 3 else None
        result_path = convert_pdf_to_docx(input_file, output_file)
        print(json.dumps({"output_path": result_path}, ensure_ascii=False))

    else:
        print(f"未知命令: {command}")
        sys.exit(1)
